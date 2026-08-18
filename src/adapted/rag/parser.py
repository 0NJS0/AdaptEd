from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from pypdf import PdfReader


class UnsupportedDocumentError(ValueError):
    pass


@dataclass
class Page:
    number: int  # 1-based
    text: str


@dataclass
class ParsedDocument:
    filename: str
    pages: list[Page] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    @property
    def page_count(self) -> int:
        return len(self.pages)


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


def parse_document(path: Path, filename: str) -> ParsedDocument:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
        )
    if ext == ".pdf":
        return _parse_pdf(path, filename)
    if ext in (".txt", ".md"):
        return _parse_text(path, filename)
    if ext == ".docx":
        return _parse_docx(path, filename)
    raise UnsupportedDocumentError(f"Unsupported file type '{ext}'.")


def _parse_pdf(path: Path, filename: str) -> ParsedDocument:
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise UnsupportedDocumentError(f"Could not read PDF: {exc}") from exc
    pages: list[Page] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            text = ""
        text = _clean(text)
        if text.strip():
            pages.append(Page(number=i + 1, text=text))
    if not pages:
        raise UnsupportedDocumentError("The PDF contains no extractable text.")
    return ParsedDocument(filename=filename, pages=pages)


def _parse_text(path: Path, filename: str) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    text = _clean(text)
    return ParsedDocument(filename=filename, pages=[Page(number=1, text=text)])


def _parse_docx(path: Path, filename: str) -> ParsedDocument:
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
    except Exception as exc:
        raise UnsupportedDocumentError(f"Could not read DOCX: {exc}") from exc
    text = _clean("\n".join(parts))
    return ParsedDocument(filename=filename, pages=[Page(number=1, text=text)])


def _clean(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
